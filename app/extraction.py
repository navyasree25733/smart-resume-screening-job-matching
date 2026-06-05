

# import fitz  # PyMuPDF
# import docx
# import io
# import asyncio

# async def extract_resume_text(file) -> str:

#     await file.seek(0)

#     content = await file.read()

#     print("Filename:", file.filename) 
#     print("File Size:", len(content))

#     await file.seek(0)

#     text = ""    
#     # Try PDF with PyMuPDF
#     try:
#         doc = fitz.open(stream=content, filetype="pdf")
#         text = ""
#         for page in doc:
#             text += page.get_text()
#         doc.close()
#         if len(text.strip()) > 50:
#             return text.strip()
#     except Exception as e:
#         print(f"PyMuPDF error: {e}")
    
#     # Try DOCX
#     try:
#         doc = docx.Document(io.BytesIO(content))
#         text = "\n".join([para.text for para in doc.paragraphs])
#         if len(text.strip()) > 50:
#             return text.strip()
#     except Exception as e:
#         print(f"DOCX error: {e}")
    
#     # Try raw bytes as fallback
#     try:
#         text = content.decode('utf-8', errors='ignore')
#         if len(text.strip()) > 50:
#             return text.strip()
#     except:
#         pass
    
#     return text if len(text.strip()) > 20 else ""


# # import fitz  # PyMuPDF
# # import docx
# # import io
# # import asyncio

# # async def extract_resume_text(file) -> str:
# #     """Extract text from PDF/DOCX with fallback"""
# #     content = await file.read()
    
# #     text = ""
    
# #     # Try PDF with PyMuPDF
# #     try:
# #         doc = fitz.open(stream=content, filetype="pdf")
# #         text = ""
# #         for page in doc:
# #             text += page.get_text()
# #         doc.close()
# #         if len(text.strip()) > 50:
# #             return text.strip()
# #     except Exception as e:
# #         print(f"PyMuPDF error: {e}")
    
# #     # Try DOCX
# #     try:
# #         doc = docx.Document(io.BytesIO(content))
# #         text = "\n".join([para.text for para in doc.paragraphs])
# #         if len(text.strip()) > 50:
# #             return text.strip()
# #     except Exception as e:
# #         print(f"DOCX error: {e}")
    
# #     # Try raw bytes as fallback
# #     try:
# #         text = content.decode('utf-8', errors='ignore')
# #         if len(text.strip()) > 50:
# #             return text.strip()
# #     except:
# #         pass
    
# #     return text if len(text.strip()) > 20 else ""


import fitz  # PyMuPDF
import docx
import io


async def extract_resume_text(file) -> str:
    """
    Extract text from PDF or DOCX resumes.

    Supports:
    - PDF (.pdf)
    - DOCX (.docx)

    Returns:
    - Extracted text
    - Empty string if extraction fails
    """

    try:

        # Reset file pointer
        await file.seek(0)

        # Read uploaded file
        content = await file.read()

        print("=" * 50)
        print("Filename:", file.filename)
        print("File Size:", len(content))
        print("=" * 50)

        if not content:
            print("ERROR: Empty file received")
            return ""

        # Reset again for later use
        await file.seek(0)

        filename = file.filename.lower()

        # =====================================
        # PDF EXTRACTION
        # =====================================

        if filename.endswith(".pdf"):

            try:

                pdf = fitz.open(
                    stream=content,
                    filetype="pdf"
                )

                text = ""

                for page in pdf:
                    text += page.get_text()

                pdf.close()

                print(
                    "PDF Text Length:",
                    len(text)
                )

                if text and text.strip():
                    return text.strip()

            except Exception as e:

                print(
                    f"PDF Extraction Error: {e}"
                )

        # =====================================
        # DOCX EXTRACTION
        # =====================================

        elif filename.endswith(".docx"):

            try:

                document = docx.Document(
                    io.BytesIO(content)
                )

                text = "\n".join(
                    para.text
                    for para in document.paragraphs
                )

                print(
                    "DOCX Text Length:",
                    len(text)
                )

                if text and text.strip():
                    return text.strip()

            except Exception as e:

                print(
                    f"DOCX Extraction Error: {e}"
                )

        # =====================================
        # FALLBACK TEXT EXTRACTION
        # =====================================

        try:

            text = content.decode(
                "utf-8",
                errors="ignore"
            )

            print(
                "Fallback Text Length:",
                len(text)
            )

            if text and text.strip():
                return text.strip()

        except Exception as e:

            print(
                f"Fallback Error: {e}"
            )

        return ""

    except Exception as e:

        print(
            f"Resume Extraction Error: {e}"
        )

        return ""

